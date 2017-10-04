#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Module for exporting cvs to Word format."""
from __future__ import print_function

import datetime
import socket
import docx
from StringIO import StringIO

from docx import Document
from docx.shared import Cm 
from docx.shared import Pt 
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_BUILTIN_STYLE
from rsum.settings.rsum import values

import home.models

CV = home.models.cv.CV


class ExportDocument(object):
    """Class to handle exporting rsum pages to Word documents.

    .. attribute:: s

       Hostname of current host.

    .. attribute:: name

       Filename to offer to the end user.
    """
    def __init__(self):
        """Initialize ExportDocument class.

        :return: None
        :rtype: None
        """
        self.s = values.get(socket.gethostname())
        s = self.s
        self.name = '{0}-cv.docx'.format(s.get('owner'))

    def export(self, cv_id):
        """Export a word document.
    
        :param cv_id: ID of CV to export.
        :type cv_id: int
        :return: Stream of Word document for end user.
        :rtype: object 
        """
        cv_instance = CV()
        cv = cv_instance.get_cv(cv_id)
        stream = StringIO()
        document = Document()
        document = self.set_styles(document)
        document = self.set_layout(document)

        for current,section in sorted(enumerate(cv.get('sections'))):
            print(section.get('name'))
            if section.get('name') == u'intro':
                document = self.add_intro(section.get('content'), document)

            if section.get('name') == u'summary':
                p = document.add_paragraph('')
                p.paragraph_format.line_spacing = 0.0
                document = self.add_summary(section.get('content'), document)

            if section.get('name') == u'skills':
                p = document.add_paragraph('')
                p.paragraph_format.line_spacing = 0.0
                document = self.add_skills(section.get('content'), document)
        
            if section.get('name') == u'experience':
                p = document.add_paragraph('')
                p.paragraph_format.line_spacing = 0.0
                p.paragraph_format.page_break_before = True
                document = self.add_experience(section.get('content'), document)
        
            if section.get('name') == u'education':
                p = document.add_paragraph('')
                p.paragraph_format.line_spacing = 0.0
                document = self.add_education(section.get('content'), document)
                
            if section.get('name') == u'contact':
                p = document.add_paragraph('')
                p.paragraph_format.line_spacing = 0.0
                document = self.add_contact(section.get('content'), document)

        document.save(stream)

        return stream 

    def add_intro(self, intro, document):
        """Add introduction section.
    
        :param intro: Introduction to add to document.
        :type intro: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Introduction.
        :rtype: object
        """
        s = self.s
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        content = {}
        for index, item in enumerate(sorted(intro)):
            print(item)
            if item.get('content')[0].get('name') == 'name':
                content.update({'name': item.get('content')[0].get('content')})
            if item.get('content')[0].get('name') == 'position':
                content.update({'position': item.get('content')[0].get('content')})
        table.cell(0,0).add_paragraph(content.get('name'), style='Heading 1')
        table.cell(0,0).add_paragraph(content.get('position'), style='Heading 2')

        table.cell(0,0).width = Cm(12)
        table.cell(0,1).add_picture('/srv/rsum/static/{0}/img/mockup/avatar-01.png'.format(s.get('template')), width=Cm(4))
        table.cell(0,1).width = Cm(4)
        table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return document 

    def add_summary(self, summary, document):
        """Add summary section.
    
        :param summary: Summary section to add to document.
        :type summary: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Summary.
        :rtype: object
        """
        s = self.s
        t = document.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, item in enumerate(sorted(summary)):
            if len(item.get('content')) == 1:
                t.cell(0,0).add_picture('/srv/rsum/static/{0}/img/500x700/01.jpg'.format(s.get('template')), width=Cm(6))
                t.cell(0,0).width = Cm(6)
                t.cell(0,1).add_paragraph('Summary', style='Heading 3')
                t.cell(0,1).add_paragraph(item.get('content')[0].get('content'), style='Normal')
                p = t.cell(0,1).paragraphs[1]
                p.paragraph_format.line_spacing = 1.0 
                t.cell(0,1).width = Cm(10)
                p = t.cell(0,0).paragraphs[0]
                p.paragraph_format.line_spacing = 0.0
                p = t.cell(0,1).paragraphs[0]
                p.paragraph_format.line_spacing = 0.0
        return document 

    def add_skills(self, skills, document):
        """Add skills section.
    
        :param skills: Skills section to add to document.
        :type summary: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Skills.
        :rtype: object
        """
        current_year = datetime.datetime.now().strftime("%Y")

        t = document.tables[1]
        t.cell(0,1).add_paragraph('Skills', style='Heading 3')
        t_sub = t.cell(0,1).add_table(rows=1, cols=2)
        t.cell(0,1).tables[0].columns[0].width = Cm(7)
        for index,item in enumerate(sorted(skills)):
            if index > 0:
                skill = ['0', '1'] 
                subskills = []
                for item_content in item.get('content'):
                    if (
                        isinstance(item_content.get('content'), unicode) and
                        item_content.get('name') == 'start'
                    ):
                        experience = int(current_year) - int(item_content.get('content'))
                        skill[1] = '{0} year(s)'.format(str(experience))
                    elif (
                        isinstance(item_content.get('content'), unicode) and
                        item_content.get('name') != 'id'
                    ):
                        name = item_content.get('content').replace('_',' ').title()
                        skill[0] = name 
                    elif (
                        isinstance(item_content.get('content'), dict)
                    ):
                        subskills.append(item_content.get('content'))
                t_sub.add_row()
                t_sub.cell(index-1,0).text = skill[0]
                p = t_sub.cell(index-1,0).paragraphs[0]
                p.style='Skill'
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = 0
                t_sub.cell(index-1,1).text = skill[1]
                p = t_sub.cell(index-1,1).paragraphs[0]
                p.style='Skill'
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = 0
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                t_sub = self.add_sub_skills(subskills, t_sub, index-1)
        return document

    def add_sub_skills(self, subs, ts, ts_index):
        """Add sub skills to skills section.
    
        :param subs: Sub skills to add to document.
        :type subs: [dict(str, str)]
        :param object ts: Table cell to update. 
        :param int ts_index: Index for current table cell. 
        :return: Document updated with sub skills.
        :rtype: object
        """
        current_year = float(datetime.datetime.now().strftime("%Y")) 
        sub_table = ts.cell(ts_index,0).add_table(rows=1,cols=2)
        for i, sub in enumerate(subs):
            name = sub.items()[0][1].get('name') 
            experience = int(current_year) - int(sub.items()[0][1].get('start'))
            experience = '{0} year(s)'.format(str(experience))
            if i == 0:
                sub_table.cell(0,0).text = name 
                sub_table.cell(0,1).text = experience
                sub_table.cell(0,0).width = Cm(5)
            else:
                sub_table.add_row()
                sub_table.cell(i,0).text = name
                sub_table.cell(i,0).width = Cm(1)
                sub_table.cell(i,1).text = experience
            p = sub_table.cell(i,0).paragraphs[0]
            p.style = 'Sub Skill'
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = 0
            p = sub_table.cell(i,1).paragraphs[0]
            p.style = 'Sub Skill'
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_after = 0
        return ts 

    def add_experience(self, experience, document):
        """Add experience section.
    
        :param [dict(str, str)] experience:
            Experience section to add to document.
        :param object document: Document to update. 
        :return: Document updated with Experience section.
        :rtype: object
        """
        s = self.s
        del experience[0]
        p = document.add_paragraph('Experience', style='Heading 3')
        t = document.add_table(rows=1, cols=3)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, value in enumerate(experience):
            if index < 3:
                pass
            elif index % 3 == 0:
                t.add_row()
            for item in value.get('content'):
                if item.get('name') == 'location':
                    location = item.get('content')
                if item.get('name') == 'duration':
                    duration = item.get('content')
                if item.get('name') == 'position':
                    position = item.get('content')
                if item.get('name') == 'projects':
                    projects = item.get('content').get('projects')
                if item.get('name') == 'id':
                    photo = item.get('content')
                if item.get('name') == 'company':
                    company = item.get('content')
            col = index % 3
            row = index / 3
            p = t.cell(row,col).paragraphs[0]
            p.paragraph_format.line_spacing = 0.0
            t.cell(row,col).add_picture('/srv/rsum/static/{0}/img/970x647/{1}.jpg'.format(s.get('template'),photo), width=Cm(4.8))
            p = t.cell(row,col).paragraphs[1]
            p.paragraph_format.space_after = 0
            p = t.cell(row,col).add_paragraph(position, style='Heading 4')
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = 0
            p = t.cell(row,col).add_paragraph(company, style='Heading 5')
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = 0
            p = t.cell(row,col).add_paragraph("{0}, {1}".format(location, duration), style='Heading 6')
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = 0
            t = self.add_projects(projects, t, row, col)
        p = document.add_paragraph('')
        p.paragraph_format.line_spacing = 0
        return document

    def add_projects(self, projects, table, row, col):
        """Add projects to experience section.
    
        :param [dict(str, str)] projects:
            Projects for a portion of Experience section. 
        :param object table: Table from current document. 
        :param int row: Index of current row.
        :param int col: Index of current col.
        :return: Updated Projects table. 
        :rtype: object
        """
        for project in projects.items():
            p = table.cell(row,col).add_paragraph(project[0], style='List Bullet')
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = 0
            actions = project[1].replace('[','').replace(']','').split("', '") 
            for action in actions:
                p = table.cell(row,col).add_paragraph(action.replace("'",""), style = 'List Bullet 2')
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = 0
        return table

    def add_education(self, education, document):
        """Add education section.
    
        :param [dict(str, str)] education:
            Education section for current document. 
        :param object document: Current document. 
        :return: Current document with Educaiton section. 
        :rtype: object
        """
        s = self.s
        p = document.add_paragraph('Education', style='Heading 3')
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = 0
        p.paragraph_format.page_break_before = True
        document.add_picture('/srv/rsum/static/{0}/img/1920x1080/01.jpg'.format(s.get('template')), width=Cm(4))
        for item in education:
            if item.get('content')[0].get('name') == u'name':
                name = item.get('content')[0].get('content')
            if item.get('content')[0].get('name') == u'duration':
                duration = item.get('content')[0].get('content')
                a = duration[:4]
                b = duration[7:]
                duration = a + ' - ' + b
            if item.get('content')[0].get('name') == u'studies':
                studies = item.get('content')[0].get('content')
            if item.get('name') == u'projects':
                projects = item.get('content')[0].get('content')
                projects = projects.replace('[','').replace(']','').split("', '")
            if item.get('name') == u'location':
                location = item.get('content')[0].get('content')
        p = document.add_paragraph(name, style='Heading 4')
        p.paragraph_format.space_before = 0
        p = document.add_paragraph(studies, style='Heading 5')
        p.paragraph_format.space_before = 0
        p = document.add_paragraph("{0}, {1}".format(location, duration), style='Heading 6')
        p.paragraph_format.space_before = 0
        for project in projects:
            p = document.add_paragraph(project.replace("'",''), style='List Bullet')
        return document

    def add_contact(self, contact, document):
        """Add contact section.
    
        :param [dict(str, str)] contact:
            Contact section for current document. 
        :param object document: Current document. 
        :return: Current document with Contact section. 
        :rtype: object
        """
        for item in contact:
            if item.get('name') == 'web':
                web = item.get('content')[0].get('content')
            if item.get('name') == 'location':
                location = item.get('content')[0].get('content')
            if item.get('name') == 'title':
                title = item.get('content')[0].get('content')
            if item.get('name') == 'message':
                message = item.get('content')[0].get('content')
            if item.get('name') == 'email':
                email = item.get('content')[0].get('content')
            if item.get('name') == 'phone':
                phone = item.get('content')[0].get('content')
        document.add_paragraph(title, style='Heading 3')
        p = document.add_paragraph(message, style='Normal')
        p.paragraph_format.space_after = 0
        t = document.add_table(rows=2,cols=6)
        p = t.cell(0,0).paragraphs[0]
        p.paragraph_format.line_spacing = 0
        t.cell(0,0).add_paragraph('Website', style='Heading 4')
        p = t.cell(0,0).add_paragraph(web, style='Heading 5')
        p.paragraph_format.space_before = 0
        p = t.cell(0,1).paragraphs[0]
        p.paragraph_format.line_spacing = 0
        t.cell(0,1).add_paragraph('Email', style='Heading 4')
        p = t.cell(0,1).add_paragraph(email, style='Heading 5')
        p.paragraph_format.space_before = 0
        p = t.cell(1,0).paragraphs[0]
        p.paragraph_format.line_spacing = 0
        t.cell(1,0).add_paragraph('Phone', style='Heading 4')
        p = t.cell(1,0).add_paragraph(phone, style='Heading 5')
        p.paragraph_format.space_before = 0
        p = t.cell(1,1).paragraphs[0]
        p.paragraph_format.line_spacing = 0
        t.cell(1,1).add_paragraph('Location', style='Heading 4')
        p = t.cell(1,1).add_paragraph(location, style='Heading 5') 
        p.paragraph_format.space_before = 0
        return document

    def set_styles(self, document):
        """Set styles in the Word document.
    
        :param object document: Current document. 
        :return: Current document with correct styles. 
        :rtype: object
        """
        style = document.styles['Heading 1'] 
        font = style.font
        font.color.rgb = RGBColor(0x51, 0x57, 0x6A) 
        font.name = 'Hind'
        font.size = Pt(24) 
        font.bold = True

        style = document.styles['Heading 2']
        font = style.font
        font.name = 'Hind'
        font.italic = False
        font.color.rgb = RGBColor(0xA6, 0xA7, 0xAA)
        font.size = Pt(16)

        style = document.styles['Heading 3']
        font = style.font
        font.name = 'Hind'
        font.color.rgb = RGBColor(0x51, 0x57, 0x6A) 
        font.size = Pt(14)
        font.bold = True

        style = document.styles['Heading 4']
        font = style.font
        font.name = 'Hind'
        font.color.rgb = RGBColor(0x51, 0x57, 0x6A) 
        font.size = Pt(8)
        font.bold = True
        font.italic = False
        
        style = document.styles['Heading 5']
        font = style.font
        font.color.rgb = RGBColor(0xA6, 0xA7, 0xAA)
        font.small_caps = True
        font.name = 'Hind'
        font.size = Pt(7)
        
        style = document.styles['Heading 6']
        font = style.font
        font.color.rgb = RGBColor(0x51, 0x57, 0x6A) 
        font.name = 'Hind'
        font.size = Pt(6)
        font.bold = True
        font.italic = False
        
        style = document.styles['List Bullet']
        font = style.font
        font.color.rgb = RGBColor(0x51, 0x57, 0x6A) 
        font.name = 'Hind'
        font.size = Pt(5)
        font.bold = True

        style = document.styles['List Bullet 2']
        font = style.font
        font.color.rgb = RGBColor(0xA6, 0xA7, 0xAA)
        font.size = Pt(5)
        font.name = 'Hind'

        document.styles.add_style('Skill', WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles['Skill']
        font = style.font
        font.name = 'Hind'
        font.color.rgb = RGBColor(0x51, 0x57, 0x6A) 
        font.size = Pt(9)
        font.bold = True

        document.styles.add_style('Sub Skill', WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles['Sub Skill']
        font = style.font
        font.name = 'Hind'
        font.color.rgb = RGBColor(0x51, 0x57, 0x6A) 
        font.size = Pt(7)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Hind'
        font.color.rgb = RGBColor(0xA6, 0xA7, 0xAA)
        font.size = Pt(11)
        return document

    def set_layout(self, document):
        """Define document layout.
    
        :param object document: Current document. 
        :return: Current document with adjusted layougt. 
        :rtype: object
        """
        sections = document.sections 
        section = sections[0] 
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.left_margin = Cm(0.5)
        section.top_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)

        return document
