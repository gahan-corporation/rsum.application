"""Intro module."""
# pylint: disable=too-few-public-methods,no-name-in-module,no-member
from django.conf import settings as django_settings
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Intro(object):
    """Class for Intro objects."""

    name = None
    document = None
    settings = django_settings

    def save(self, section, document, graphics):
        """Add introduction section.

        :param intro: Introduction to add to document.
        :type intro: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Introduction.
        :rtype: object
        """
        if graphics:
            document = self.get_intro_graphics(section, document)
        else:
            document = self.get_intro(section, document)
        return document

    def get_intro(self, section, document):
        """Add introduction section.

        :param intro: Introduction to add to document.
        :type intro: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Introduction.
        :rtype: object
        """
        intro = section
        self.document = document
        paragraph = document.add_paragraph(
            intro.get('name'), style='Heading 1')

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        paragraph = document.add_paragraph(
            intro.get('position'), style='Heading 2')

        return document

    def get_intro_graphics(self, section, document):
        """Add introduction section.
        :param intro: Introduction to add to document.
        :type intro: [dict(str, str)]
        :param document: Current document.
        :type document: object
        :return: Document updated with Introduction.
        :rtype: object
        """
        intro = section
        settings = self.settings
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).width = Cm(12)

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 0).add_paragraph(
            intro.get('name'),
            style='Heading 1'
        )
        table.cell(0, 0).add_paragraph(
            intro.get('position'),
            style='Heading 2')

        table.cell(0, 1).width = Cm(4)
        table.cell(0, 1).add_picture(
            'static/profiles/{0}/img/mockup/avatar-02.png'.format(
                settings.DIR))
        table.cell(
            0,
            1
        ).paragraphs[
            0
        ].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return document
